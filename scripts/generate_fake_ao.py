#!/usr/bin/env python3
"""Génère un faux dossier AO avec de vrais fichiers PDF/DOCX/XLSX lisibles."""

import io
import shutil
from pathlib import Path
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment


OUT = Path("fake_ao_test2")


def make_pdf(path: Path, title: str, lines: list[str]):
    """Crée un PDF minimal valide lisible par pypdf."""
    path.parent.mkdir(parents=True, exist_ok=True)

    # Contenu de la page (opérateurs PDF)
    content_lines = [f"BT", f"/F1 14 Tf", f"50 750 Td", f"({title}) Tj"]
    y = 720
    for line in lines:
        # Nettoyer les caractères non-ASCII pour PDF Type1
        safe = line.encode("ascii", "replace").decode("ascii")
        content_lines.append(f"0 -18 Td ({safe}) Tj")
    content_lines.append("ET")
    stream = "\n".join(content_lines).encode("latin-1", "replace")

    # Construction manuelle du PDF avec offsets corrects
    objects = []

    obj1 = b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n"
    obj2 = b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n"
    obj3 = (
        b"3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792]\n"
        b"   /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>\nendobj\n"
    )
    obj4 = (
        b"4 0 obj\n<< /Length " + str(len(stream)).encode() + b" >>\nstream\n"
        + stream + b"\nendstream\nendobj\n"
    )
    obj5 = b"5 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n"

    header = b"%PDF-1.4\n"
    body = header
    offsets = []
    for obj in [obj1, obj2, obj3, obj4, obj5]:
        offsets.append(len(body))
        body += obj

    xref_offset = len(body)
    xref = b"xref\n0 6\n0000000000 65535 f \n"
    for off in offsets:
        xref += f"{off:010d} 00000 n \n".encode()

    trailer = (
        xref
        + b"trailer\n<< /Size 6 /Root 1 0 R >>\n"
        + b"startxref\n" + str(xref_offset).encode() + b"\n%%EOF\n"
    )

    path.write_bytes(body + trailer)


def make_docx(path: Path, title: str, sections: dict[str, str]):
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading(title, 0)
    for heading, body in sections.items():
        doc.add_heading(heading, 1)
        doc.add_paragraph(body)
    doc.save(path)


def make_xlsx(path: Path, title: str, headers: list[str], rows: list[list]):
    path.parent.mkdir(parents=True, exist_ok=True)
    wb = Workbook()
    ws = wb.active
    ws.title = title[:31]
    ws.append([title])
    ws.merge_cells(f"A1:{chr(64+len(headers))}1")
    ws["A1"].font = Font(bold=True, size=14)
    ws["A1"].alignment = Alignment(horizontal="center")
    ws.append([])
    ws.append(headers)
    header_row = ws[3]
    fill = PatternFill("solid", fgColor="4472C4")
    for cell in header_row:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = fill
    for row in rows:
        ws.append(row)
    for i in range(1, len(headers) + 1):
        ws.column_dimensions[ws.cell(row=3, column=i).column_letter].width = 22
    wb.save(path)


# ── FICHIERS RACINE (désorganisés, noms génériques) ──────────────────────────

make_pdf(OUT / "Document1.pdf", "CCTP - Lot 01 Gros Oeuvre", [
    "CAHIER DES CLAUSES TECHNIQUES PARTICULIERES",
    "Marche de travaux - Lot 01 Gros Oeuvre",
    "Operation : Construction d'un immeuble de logements collectifs",
    "Maitre d'ouvrage : Commune de Marly-le-Roi (78)",
    "Le present CCTP definit les prescriptions techniques applicables aux travaux de gros oeuvre.",
    "Art. 1 - Generalites : L'entrepreneur doit respecter les DTU en vigueur.",
    "Art. 2 - Terrassements : Fouilles en rigoles, profondeur -1.80m/TN.",
    "Art. 3 - Fondations : Semelles filantes beton C25/30, armatures HA.",
    "Art. 4 - Voiles beton : Ep. 20cm, beton C30/37, coffrage soigne.",
])

make_pdf(OUT / "scan0042.pdf", "Acte d'Engagement - Lot 02 Charpente", [
    "ACTE D'ENGAGEMENT",
    "Lot 02 - Charpente Couverture Etancheite",
    "Je soussigne, representant de SMA BTP, m'engage a realiser les travaux",
    "decrits au prix forfaitaire toutes taxes comprises.",
    "Montant HT : 245 000 EUR",
    "TVA 20% : 49 000 EUR",
    "Montant TTC : 294 000 EUR",
    "Validite de l'offre : 120 jours",
    "Date et signature : Paris, le 10/02/2025",
])

make_pdf(OUT / "Nouveau document.pdf", "Reglement de Consultation", [
    "REGLEMENT DE CONSULTATION",
    "Procedure : Appel d'offres ouvert",
    "Objet : Travaux de construction d'une ecole primaire",
    "Commune de Herblay-sur-Seine (95)",
    "Date limite de remise des offres : 15/03/2025 a 12h00",
    "Criteres de jugement : Prix 60% / Valeur technique 40%",
    "Cautionnement : 5% du montant du marche",
    "Duree des travaux : 18 mois",
])

make_docx(OUT / "copie finale v3.docx", "Rapport Initial de Controle Technique (RICT)", {
    "Identification de l'operation": (
        "Operation : Residence Les Acacias - Marly-le-Roi\n"
        "Bureau de controle : SOCOTEC\n"
        "Phase : PRO"
    ),
    "Donnees sismiques et geotechniques": (
        "Indice sismique : Zone 1 (faible sismicite)\n"
        "Classe d'exposition : XC2\n"
        "Etude de sol G2 PRO disponible."
    ),
    "Avis technique": (
        "La conception respecte les regles parasismiques PS-MI.\n"
        "Aucune reserve majeure identifiee a ce stade."
    ),
})

make_pdf(OUT / "CCAP lot 3.pdf", "CCAP - Lot 03 Plomberie Sanitaires", [
    "CAHIER DES CLAUSES ADMINISTRATIVES PARTICULIERES",
    "Lot 03 - Plomberie Sanitaires",
    "Delai d'execution : 8 mois",
    "Modalites de reglement : Acomptes mensuels sur situations de travaux",
    "Retenue de garantie : 5%",
    "Assurance decennale obligatoire.",
    "Penalites de retard : 1/3000e du montant HT par jour calendaire.",
])

make_pdf(OUT / "image001.pdf", "DC2 - Declaration du Candidat", [
    "DC2 - DECLARATION DU CANDIDAT INDIVIDUEL",
    "Identification du candidat : SMA BTP",
    "SIREN : 775 688 732",
    "Forme juridique : Societe anonyme",
    "Chiffre d'affaires HT N-1 : 1 245 000 000 EUR",
    "Chiffre d'affaires HT N-2 : 1 189 000 000 EUR",
    "Chiffre d'affaires HT N-3 : 1 134 000 000 EUR",
    "Effectif moyen annuel : 1 850 salaries",
])

make_pdf(OUT / "questions reponses.pdf", "Questions-Reponses N°2", [
    "QUESTIONS - REPONSES N2",
    "Consultation : Construction ecole primaire Herblay",
    "Q1 : Les fondations sur pieux sont-elles incluses dans le lot GO ?",
    "R1 : Oui, les fondations speciales sont a la charge du lot 01.",
    "Q2 : Le DPGF doit-il etre remis signe ?",
    "R2 : Oui, le DPGF signe est une piece obligatoire de l'offre.",
    "Q3 : Le délai de 18 mois inclut-il le repli de chantier ?",
    "R3 : Oui, le délai comprend l'installation et le repli.",
])

make_pdf(OUT / "Attestation 2024.pdf", "Attestation de Regularite Fiscale 2024", [
    "ATTESTATION DE REGULARITE FISCALE",
    "Delivree par la Direction Generale des Finances Publiques",
    "Au profit de : SMA BTP - SIREN 775 688 732",
    "Valable jusqu'au : 31/12/2024",
    "L'entreprise est a jour de ses obligations fiscales.",
    "Reference : DGFiP-2024-AF-0042187",
])

make_pdf(OUT / "plan masse.pdf", "Plan de Masse - Ind. C", [
    "PLAN DE MASSE",
    "Echelle 1/500",
    "Operation : Construction residence Herblay-sur-Seine",
    "Architecte : Cabinet DURAND & Associes",
    "Indice : C - Date : janvier 2025",
    "Surface terrain : 2 450 m2",
    "SHON : 1 820 m2",
    "Emprise au sol : 412 m2",
    "Coefficient d'emprise : 16.8%",
])

make_pdf(OUT / "urssaf.pdf", "Attestation de Vigilance URSSAF", [
    "ATTESTATION DE VIGILANCE",
    "Delivree par l'URSSAF Ile-de-France",
    "Entreprise : SMA BTP",
    "SIRET : 775 688 732 00042",
    "Periode : 4eme trimestre 2024",
    "L'entreprise est a jour de ses cotisations sociales.",
    "Document valable 6 mois a compter de sa date d'emission.",
    "Code de verification : UV-2024-IDF-887654",
])

make_pdf(OUT / "Extrait KBIS.pdf", "Extrait KBIS - SMA BTP", [
    "EXTRAIT KBIS",
    "Registre du Commerce et des Societes de Paris",
    "Denomination : SMA BTP",
    "Siege social : 214 boulevard Saint-Germain - 75007 Paris",
    "Capital social : 229 090 160 EUR",
    "Date d'immatriculation : 12/04/1907",
    "Code APE : 6512Z - Autres assurances",
    "Directeur General : M. Jean-Pierre MARTIN",
])

make_pdf(OUT / "PGCSPS_v2_final.pdf", "Plan General de Coordination SPS", [
    "PLAN GENERAL DE COORDINATION SPS",
    "Operation : Construction immeuble logements - Herblay",
    "Coordonnateur SPS : Cabinet SPS Consulting",
    "Phase : PRO/EXE",
    "Effectif maximal previsible : 45 compagnons",
    "Duree des travaux : 18 mois",
    "Risques particuliers identifies :",
    "- Travaux en hauteur (charpente R+4)",
    "- Coactivite GO / Second oeuvre a partir du 9eme mois",
    "- Proximite riverains : cloture de chantier obligatoire",
])

make_pdf(OUT / "Rapport amiante - bat A.pdf", "Diagnostic Amiante Avant Travaux - Bat. A", [
    "RAPPORT DE DIAGNOSTIC AMIANTE AVANT TRAVAUX (DAT)",
    "Batiment A - Residence Les Tilleuls",
    "Donneur d'ordre : SMA BTP",
    "Operateur de reperage : DEKRA Industrial",
    "Certification : COFRAC n°1-1234",
    "Materiaux contenant de l'amiante detectes :",
    "- Dalles de sol vinyle (couche adhesive) : PRESENCE CONFIRMEE",
    "- Flocage des poutres en beton : Absence confirmee",
    "- Calorifugeage des canalisations : Absence confirmee",
    "Recommandation : retrait dalles avant demolition, entreprise certifiee.",
])

make_pdf(OUT / "etude sol G2 PRO.pdf", "Etude Geotechnique G2 PRO", [
    "ETUDE GEOTECHNIQUE DE PROJET (G2 PRO)",
    "Site : Rue des Lilas - Herblay-sur-Seine (95)",
    "Maitre d'ouvrage : Commune de Herblay",
    "Geotechnicien : GEOTECH Ile-de-France",
    "Sondages realises : 6 (SC1 a SC6), profondeur max 12m",
    "Nappe phreatique : presence a -3.50m / TN",
    "Contrainte admissible des sols : 180 kPa a 1.50m de profondeur",
    "Fondations recommandees : semelles filantes, ancrage 1.80m / TN",
    "Risque sismique : Zone 1 - faible",
])

make_pdf(OUT / "convention DO.pdf", "Convention Assurance Dommages-Ouvrage", [
    "CONVENTION D'ASSURANCE DOMMAGES-OUVRAGE",
    "Souscripteur : Commune de Marly-le-Roi",
    "Assureur : SMA BTP",
    "N° de police : DO-2025-078-341",
    "Operation : Construction immeuble 24 logements",
    "Valeur declaree a l'achevement : 4 200 000 EUR HT",
    "Franchise contractuelle : 3 000 EUR",
    "Duree de garantie : 10 ans a compter de la reception",
    "Prise d'effet : date d'ouverture de chantier",
])

make_pdf(OUT / "AE TRC signe.pdf", "Acte d'Engagement TRC - Signe", [
    "ACTE D'ENGAGEMENT - TOUS RISQUES CHANTIER",
    "Souscripteur : Commune d'Herblay-sur-Seine",
    "Assureur : SMA BTP",
    "N° de police : TRC-2025-095-112",
    "Montant des travaux assures : 6 800 000 EUR HT",
    "Duree de chantier : 20 mois",
    "Date de debut previsionnelle : 01/03/2025",
    "Signe le : 14/02/2025 - Cachet et signature SMA BTP",
])

make_pdf(OUT / "DC1 pouvoir.pdf", "DC1 - Lettre de Candidature + Pouvoir", [
    "DC1 - LETTRE DE CANDIDATURE",
    "Identification du candidat : SMA BTP",
    "Represente par : M. Jean-Pierre MARTIN, Directeur General Delegue",
    "Pouvoir joint : Oui (voir piece separee)",
    "Marche : Construction ecole primaire - Commune de Herblay-sur-Seine",
    "Lot : Lot unique - TCE",
    "Date et signature : Paris, le 10/02/2025",
])

make_pdf(OUT / "NoticeAccessibilite.pdf", "Notice Accessibilite PMR", [
    "NOTICE D'ACCESSIBILITE PMR",
    "Operation : Construction residence 24 logements",
    "Maitre d'oeuvre : Cabinet DURAND & Associes",
    "Referentiel : Arrete du 24 decembre 2015",
    "Ascenseur : obligatoire (R+3) - Largeur cabine min. 1.10m",
    "Logements adaptables : 100% des logements",
    "Logements accessibles : 20% (5 logements sur 24)",
    "Cheminement exterieur accessible : conforme",
    "Stationnement PMR : 2 places reservees",
])

make_pdf(OUT / "Memoire_technique_v1_FINAL_vf.pdf", "Memoire Technique de Gestion SMA BTP", [
    "MEMOIRE TECHNIQUE DE GESTION",
    "SMA BTP - Assurance Construction",
    "1. Organisation de la gestion des sinistres",
    "Notre dispositif repose sur 3 niveaux d'intervention :",
    "- Niveau 1 : gestionnaire dossier (sinistres < 50k EUR)",
    "- Niveau 2 : expert mandataire (50k - 500k EUR)",
    "- Niveau 3 : cellule grands sinistres (> 500k EUR)",
    "Delai moyen de traitement : 45 jours",
    "Reseau d'experts agrees : 320 experts sur toute la France",
    "Certification ISO 9001 : oui, depuis 2008",
])

make_pdf(OUT / "Pouvoir_signature_Martin.pdf", "Delegation de Pouvoir - M. MARTIN", [
    "DELEGATION DE POUVOIR",
    "SMA BTP - Societe anonyme au capital de 229 090 160 EUR",
    "M. Jean-Pierre MARTIN est habilite a signer tout acte d'engagement",
    "dans le cadre des marches publics d'assurance construction,",
    "sans limitation de montant.",
    "Fait a Paris le 01/01/2025",
    "Le President Directeur General : signature + cachet SMA BTP",
])

make_pdf(OUT / "Copie_permis_construire_PC025.pdf", "Arrete Permis de Construire PC025", [
    "ARRETE DE PERMIS DE CONSTRUIRE",
    "Commune de Herblay-sur-Seine - Service Urbanisme",
    "N° de dossier : PC 095 306 24 W 0042",
    "Petitionnaire : SCI Les Lilas",
    "Objet : Construction d'un groupe scolaire 12 classes",
    "Surface de plancher : 2 340 m2",
    "Hauteur maximale : 9.50m / TN",
    "Delivre le : 20/09/2024",
    "Le Maire de Herblay-sur-Seine - recours possible sous 2 mois",
])

make_xlsx(OUT / "planning.xlsx",
    "Planning TCE - Residence Les Acacias",
    ["Tache", "Debut", "Fin", "Duree (j)", "Responsable", "Avancement"],
    [
        ["Installation chantier", "01/03/2025", "15/03/2025", 15, "GO", "0%"],
        ["Terrassements / VRD", "16/03/2025", "31/03/2025", 16, "GO", "0%"],
        ["Fondations", "01/04/2025", "30/04/2025", 30, "GO", "0%"],
        ["Gros oeuvre RDC", "01/05/2025", "30/06/2025", 61, "GO", "0%"],
        ["Gros oeuvre R+1", "01/07/2025", "31/08/2025", 62, "GO", "0%"],
        ["Gros oeuvre R+2", "01/09/2025", "30/09/2025", 30, "GO", "0%"],
        ["Gros oeuvre R+3", "01/10/2025", "31/10/2025", 31, "GO", "0%"],
        ["Charpente couverture", "01/11/2025", "30/11/2025", 30, "Charpente", "0%"],
        ["Etancheite toiture", "01/12/2025", "15/12/2025", 15, "Etancheite", "0%"],
        ["Menuiseries ext.", "16/12/2025", "31/01/2026", 47, "Menuiseries", "0%"],
        ["Second oeuvre", "01/02/2026", "30/04/2026", 88, "Tous corps", "0%"],
        ["VRD definitifs", "01/05/2026", "31/05/2026", 31, "VRD", "0%"],
        ["Reception", "15/06/2026", "15/06/2026", 1, "MOA/MOE", "0%"],
    ]
)

# ── SOUS-DOSSIER : Divers ─────────────────────────────────────────────────────

make_pdf(OUT / "Divers" / "RIB_SMABTP_2024.pdf", "RIB - SMA BTP", [
    "RELEVE D'IDENTITE BANCAIRE",
    "Titulaire : SMA BTP",
    "Banque : Credit Agricole CIB",
    "IBAN : FR76 1820 6004 4132 4567 8901 234",
    "BIC : AGRIFRPP882",
    "Domiciliation : Credit Agricole CIB - Paris",
])

make_pdf(OUT / "Divers" / "Extrait_Banque_France.pdf", "Avis Banque de France - SMA BTP", [
    "AVIS DE LA BANQUE DE FRANCE",
    "Cotation APCR : 3+ (capacite de remboursement satisfaisante)",
    "Entreprise : SMA BTP - SIREN 775 688 732",
    "Date d'emission : 15/01/2025",
    "Document confidentiel - Usage exclusif marches publics",
])

make_pdf(OUT / "Divers" / "CCTP lot 4 electricite.pdf", "CCTP - Lot 04 Electricite", [
    "CAHIER DES CLAUSES TECHNIQUES PARTICULIERES",
    "Lot 04 - Electricite Courants Forts et Faibles",
    "Norme applicable : NF C 15-100 (edition 2002 + amendements)",
    "Puissance souscrite previsionnelle : 3 x 400A",
    "Eclairage de securite : BAES autonomes 1h",
    "GTL par logement obligatoire",
    "Tableau general BT : armoire metallique IP55",
    "VMC : commande automatique par sonde CO2",
])

# Doublon intentionnel (même contenu, nommé différemment)
shutil.copy(OUT / "Extrait KBIS.pdf", OUT / "Divers" / "KBIS_copie_backup.pdf")
shutil.copy(OUT / "urssaf.pdf", OUT / "URSSAF_attestation_copie.pdf")

# ── SOUS-DOSSIER : Plans architecte ──────────────────────────────────────────

make_pdf(OUT / "Plans architecte" / "Plan_RDC_indA.pdf", "Plan RDC - Ind. A", [
    "PLAN ARCHITECTURAL RDC",
    "Echelle 1/100 - Indice A",
    "Operation : Residence Les Acacias - Marly-le-Roi",
    "Architecte : Cabinet DURAND & Associes",
    "Surface RDC : 412 m2 SHON",
    "Composition : Hall d'entree 28m2 + local velos + local poubelles + 4 logements",
    "Logement T2 : 2 unites de 48m2",
    "Logement T3 : 2 unites de 68m2",
])

make_pdf(OUT / "Plans architecte" / "coupe AA.pdf", "Coupe AA - Longitudinale", [
    "COUPE AA - LONGITUDINALE",
    "Echelle 1/100",
    "Hauteur sous plafond : 2.50m (RDC) / 2.50m (etages courants)",
    "Hauteur totale batiment : 14.20m / TN",
    "Nombre de niveaux : RDC + R+3",
    "Acrotere : 0.60m",
    "Sous-sol : non (vide sanitaire 0.80m)",
])

make_pdf(OUT / "Plans architecte" / "facade_nord.pdf", "Elevation Facade Nord - Ind. B", [
    "ELEVATION FACADE NORD",
    "Echelle 1/100 - Indice B",
    "Materiaux : Enduit gratte blanc, menuiseries aluminium gris anthracite RAL 7016",
    "Occultations : Volets roulants electriques motorises",
    "Bardage bois : Pin douglas traite classe 4",
    "Balcons : Garde-corps acier galvanise + verre feuillete 8.8.2",
])

# ── SOUS-DOSSIER : Contrat MOE ────────────────────────────────────────────────

make_pdf(OUT / "Contrat MOE" / "AE_MOE_signe.pdf", "Acte d'Engagement MOE - Signe", [
    "ACTE D'ENGAGEMENT MAITRISE D'OEUVRE",
    "Maitre d'ouvrage : Commune de Marly-le-Roi",
    "Maitre d'oeuvre : Cabinet DURAND & Associes SARL",
    "Mission complète : DIAG / PRO / DCE / ACT / DET / AOR",
    "Taux d'honoraires : 9.2% du montant HT des travaux",
    "Base de calcul previsionnelle : 4 200 000 EUR HT",
    "Honoraires previsionnels : 386 400 EUR HT",
    "Signe le 05/11/2024 - Cachet + signature des deux parties",
])

make_docx(OUT / "Contrat MOE" / "CCAP_MOE.docx", "CCAP - Contrat de Maitrise d'Oeuvre", {
    "Objet du contrat": (
        "Maitrise d'oeuvre complete pour la construction d'une residence de 24 logements.\n"
        "Commune de Marly-le-Roi (78)."
    ),
    "Duree de mission": (
        "36 mois a compter de l'ordre de service de demarrage.\n"
        "Phases : PRO (6 mois), DCE/ACT (3 mois), DET (18 mois), AOR (6 mois)."
    ),
    "Assurances": (
        "Assurance RC professionnelle obligatoire : minimum 2 000 000 EUR par sinistre.\n"
        "Justificatif a fournir avant tout commencement de mission."
    ),
    "Sous-traitance": (
        "Toute sous-traitance est soumise a accord prealable ecrit du maitre d'ouvrage.\n"
        "Le mandataire reste seul responsable vis-a-vis du maitre d'ouvrage."
    ),
})

# ── SOUS-DOSSIER : old (obsoletes) ───────────────────────────────────────────

make_pdf(OUT / "old" / "AE_v1_obsolete.pdf", "AE Lot 02 - VERSION OBSOLETE v1", [
    "ACTE D'ENGAGEMENT VERSION 1 - OBSOLETE",
    "NE PAS UTILISER - Remplace par AE lot 02 version finale",
    "Montant initial : 230 000 EUR HT (revise a 245 000 EUR HT)",
    "Motif de revision : reevaluation quantites charpente (avenant n°1)",
])

make_pdf(OUT / "old" / "questions v1.pdf", "Q&R N°1 - Supersede", [
    "QUESTIONS REPONSES N°1 - SUPERSEDE PAR VERSION 2",
    "Q1 : Quelles sont les assurances obligatoires ?",
    "R1 : RC decennale, dommages aux existants, TRC.",
    "Q2 : Les variantes sont-elles autorisees ?",
    "R2 : Non, aucune variante n'est autorisee pour ce marche.",
])

# ── SOUS-DOSSIER : Administratif (doublons / mal classes) ────────────────────

make_pdf(OUT / "Administratif" / "URSSAF_attestation_copie.pdf", "Attestation URSSAF - Copie", [
    "ATTESTATION DE VIGILANCE - COPIE",
    "Delivree par l'URSSAF Ile-de-France",
    "Entreprise : SMA BTP - SIRET : 775 688 732 00042",
    "Periode : 4eme trimestre 2024",
    "CODE DE VERIFICATION : UV-2024-IDF-887654",
    "(Copie conforme a l'original - a ne pas utiliser pour soumission)",
])

make_xlsx(OUT / "Administratif" / "Budget_previsionnel_TCE.xlsx",
    "Budget Previsionnel - Tous Corps d'Etat",
    ["Lot", "Intitule", "Montant HT (EUR)", "% du total", "Titulaire pressenti"],
    [
        ["Lot 01", "Gros Oeuvre", 1_850_000, "27%", "A definir"],
        ["Lot 02", "Charpente Couverture Etancheite", 420_000, "6%", "A definir"],
        ["Lot 03", "Platrerie Isolation", 380_000, "6%", "A definir"],
        ["Lot 04", "Menuiseries Exterieures", 290_000, "4%", "A definir"],
        ["Lot 05", "Menuiseries Interieures", 180_000, "3%", "A definir"],
        ["Lot 06", "Revetements de sols", 220_000, "3%", "A definir"],
        ["Lot 07", "Peinture", 160_000, "2%", "A definir"],
        ["Lot 08", "Plomberie Sanitaires", 310_000, "5%", "A definir"],
        ["Lot 09", "Chauffage VMC", 280_000, "4%", "A definir"],
        ["Lot 10", "Electricite CFO/CFA", 340_000, "5%", "A definir"],
        ["Lot 11", "Ascenseur", 85_000, "1%", "A definir"],
        ["Lot 12", "VRD Espaces verts", 195_000, "3%", "A definir"],
        ["", "TOTAL HT", 4_710_000, "100%", ""],
        ["", "Provision imprevus 8%", 376_800, "", ""],
        ["", "TOTAL AVEC IMPREVU", 5_086_800, "", ""],
    ]
)

print(f"Dossier genere : {OUT}/")
print(f"Fichiers : {sum(1 for _ in OUT.rglob('*') if _.is_file())} fichiers")
for f in sorted(OUT.rglob("*")):
    if f.is_file():
        print(f"  {f.relative_to(OUT)}")
